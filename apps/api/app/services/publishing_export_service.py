# apps/api/app/services/publishing_export_service.py
#
# Generates downloadable files from story data — PDF, EPUB, DOCX, and a
# submission zip. Each function returns (bytes, filename, mimetype).
#
# Quality notes:
#   • PDF is laid out as a real 6×9" trade book (Noto Serif body, embedded), with a
#     title page, copyright page, running heads, page numbers, and a navigable
#     outline. Non-Latin text (Arabic, Cyrillic, accents, smart punctuation) renders
#     correctly because we embed Unicode TTFs instead of ReportLab's Latin-1 cores.
#   • Arabic is letter-joined (arabic-reshaper) and bidi-reordered (python-bidi) and
#     set right-to-left in Noto Naskh Arabic — ReportLab does not shape Arabic itself.
#   • **bold** / *italic* markdown is rendered as real emphasis (not stripped).
# Deps: reportlab>=4, ebooklib>=0.18, python-docx>=1.1, Pillow>=10,
#       arabic-reshaper>=3, python-bidi>=0.4

from __future__ import annotations

import html as _html
import io
import os
import re
import zipfile
from datetime import datetime, timezone
from typing import Optional

# ── PDF ─────────────────────────────────────────────────────────────────────
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_RIGHT
from reportlab.platypus import (
    BaseDocTemplate, Paragraph, Spacer, PageBreak, Image, Frame, PageTemplate,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ── EPUB ─────────────────────────────────────────────────────────────────────
from ebooklib import epub

# ── DOCX ─────────────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Fonts — embed Unicode TTFs so non-Latin text doesn't render as boxes.
# ---------------------------------------------------------------------------

_FONTS_DIR = os.path.normpath(os.path.join(os.path.dirname(__file__), "..", "assets", "fonts"))


def _register_fonts() -> bool:
    """Register the bundled Noto fonts with ReportLab. Returns False (→ fall back
    to the Latin-1 core fonts) if the TTFs are somehow missing, so export never
    hard-fails on a packaging slip."""
    try:
        reg = pdfmetrics.registerFont
        reg(TTFont("BookSerif", os.path.join(_FONTS_DIR, "NotoSerif-Regular.ttf")))
        reg(TTFont("BookSerif-Bold", os.path.join(_FONTS_DIR, "NotoSerif-Bold.ttf")))
        reg(TTFont("BookSerif-Italic", os.path.join(_FONTS_DIR, "NotoSerif-Italic.ttf")))
        reg(TTFont("BookSerif-BoldItalic", os.path.join(_FONTS_DIR, "NotoSerif-BoldItalic.ttf")))
        pdfmetrics.registerFontFamily(
            "BookSerif", normal="BookSerif", bold="BookSerif-Bold",
            italic="BookSerif-Italic", boldItalic="BookSerif-BoldItalic",
        )
        reg(TTFont("BookArabic", os.path.join(_FONTS_DIR, "NotoNaskhArabic-Regular.ttf")))
        reg(TTFont("BookArabic-Bold", os.path.join(_FONTS_DIR, "NotoNaskhArabic-Bold.ttf")))
        pdfmetrics.registerFontFamily(
            "BookArabic", normal="BookArabic", bold="BookArabic-Bold",
            italic="BookArabic", boldItalic="BookArabic-Bold",
        )
        return True
    except Exception:
        return False


_FONTS_OK = _register_fonts()
SERIF        = "BookSerif" if _FONTS_OK else "Times-Roman"
SERIF_BOLD   = "BookSerif-Bold" if _FONTS_OK else "Times-Bold"
SERIF_ITALIC = "BookSerif-Italic" if _FONTS_OK else "Times-Italic"
ARABIC       = "BookArabic" if _FONTS_OK else "Times-Roman"
ARABIC_BOLD  = "BookArabic-Bold" if _FONTS_OK else "Times-Bold"

# ── Arabic shaping (ReportLab doesn't join/bidi Arabic itself) ───────────────
try:
    import arabic_reshaper
    from bidi.algorithm import get_display
    _HAS_SHAPING = True
except Exception:  # pragma: no cover
    _HAS_SHAPING = False

_RTL_RE = re.compile(r"[؀-ۿݐ-ݿࢠ-ࣿﭐ-﷿ﹰ-﻿֐-׿]")


def _is_rtl(text: str) -> bool:
    return bool(_RTL_RE.search(text or ""))


def _shape(text: str) -> str:
    """Letter-join + bidi-reorder so Arabic displays correctly in ReportLab."""
    if not _HAS_SHAPING:
        return text
    try:
        return get_display(arabic_reshaper.reshape(text))
    except Exception:
        return text


_AR_DIGITS = str.maketrans("0123456789", "٠١٢٣٤٥٦٧٨٩")


def _ar_num(n) -> str:
    """Western digits → Arabic-Indic numerals (e.g. 12 → ١٢)."""
    return str(n).translate(_AR_DIGITS)


# ---------------------------------------------------------------------------
# Text helpers
# ---------------------------------------------------------------------------

def _word_count(text: str) -> int:
    return len(re.findall(r"\w+", text))


def _clean_text(text: str) -> str:
    """Strip markdown markers — for titles/headings and plain-text contexts."""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"^#+ ", "", text, flags=re.M)
    return text.strip()


def _paragraphs(content: str) -> list[str]:
    return [p.strip() for p in (content or "").split("\n\n") if p.strip()]


def _rl_markup(text: str) -> str:
    """Escape prose for a ReportLab Paragraph, then render **bold**/*italic*/_italic_
    as <b>/<i>. Escaping first makes a literal '<' or '&' inert, so the only tags in
    the result are the emphasis ones we add."""
    t = _html.escape(text, quote=False)
    t = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", t, flags=re.S)
    t = re.sub(r"(?<![\w*])\*(?!\s)([^*\n]+?)(?<!\s)\*(?![\w*])", r"<i>\1</i>", t)
    t = re.sub(r"(?<![\w_])_(?!\s)([^_\n]+?)(?<!\s)_(?![\w_])", r"<i>\1</i>", t)
    return t


def _epub_markup(text: str) -> str:
    """Same as _rl_markup but emits semantic <strong>/<em> for XHTML."""
    t = _html.escape(_clean_text_keep_marks(text), quote=False)
    t = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", t, flags=re.S)
    t = re.sub(r"(?<![\w*])\*(?!\s)([^*\n]+?)(?<!\s)\*(?![\w*])", r"<em>\1</em>", t)
    t = re.sub(r"(?<![\w_])_(?!\s)([^_\n]+?)(?<!\s)_(?![\w_])", r"<em>\1</em>", t)
    return t


def _clean_text_keep_marks(text: str) -> str:
    """Drop heading hashes but keep * / _ so emphasis can be rendered."""
    return re.sub(r"^#+ ", "", (text or ""), flags=re.M).strip()


# ---------------------------------------------------------------------------
# PDF Export — 6×9" trade book
# ---------------------------------------------------------------------------

class _BookDoc(BaseDocTemplate):
    """BaseDocTemplate that emits a navigable PDF outline from chapter titles."""
    def afterFlowable(self, flowable):
        if isinstance(flowable, Paragraph) and flowable.style.name == "ChapterTitle":
            label = getattr(flowable, "_outline_label", None) or flowable.getPlainText()
            n = getattr(self, "_outline_n", 0) + 1
            self._outline_n = n
            key = f"chap-{n}"
            self.canv.bookmarkPage(key)
            self.canv.addOutlineEntry(label, key, level=0, closed=False)


def export_pdf(
    title: str,
    author: str,
    tagline: Optional[str],
    chapters: list[dict],  # [{"number": int, "title": str, "content": str}]
    cover_bytes: Optional[bytes] = None,
) -> tuple[bytes, str, str]:
    """Returns (bytes, filename, mimetype). A 6×9 trade-paperback PDF."""
    buf = io.BytesIO()
    PW, PH = 6 * inch, 9 * inch

    doc = _BookDoc(
        buf, pagesize=(PW, PH),
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.8 * inch, bottomMargin=0.75 * inch,
        title=_clean_text(title), author=author,  # → PDF metadata
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="main")

    title_rtl = _is_rtl(title)
    running = _shape(_clean_text(title)) if title_rtl else _clean_text(title)
    head_font = ARABIC if title_rtl else SERIF_ITALIC

    def _deco(canvas, d):
        # Front matter (title + copyright = pages 1-2) gets no running head/number.
        if d.page <= 2:
            return
        canvas.saveState()
        canvas.setFillColor(colors.HexColor("#9a9a9a"))
        canvas.setFont(SERIF, 9)
        canvas.drawCentredString(PW / 2, 0.5 * inch, _ar_num(d.page) if title_rtl else str(d.page))
        canvas.setFont(head_font, 8.5)
        canvas.drawCentredString(PW / 2, PH - 0.55 * inch, running)
        canvas.restoreState()

    doc.addPageTemplates([PageTemplate(id="main", frames=frame, onPage=_deco)])

    base = getSampleStyleSheet()["Normal"]
    title_style = ParagraphStyle("Title", parent=base, fontName=(ARABIC_BOLD if title_rtl else SERIF_BOLD),
                                 fontSize=30, leading=38, alignment=TA_CENTER, spaceAfter=10)
    author_style = ParagraphStyle("Author", parent=base, fontName=(ARABIC if title_rtl else SERIF_ITALIC),
                                  fontSize=14, leading=20, alignment=TA_CENTER, textColor=colors.HexColor("#444444"))
    tagline_style = ParagraphStyle("Tagline", parent=base, fontName=(ARABIC if _is_rtl(tagline or "") else SERIF_ITALIC),
                                   fontSize=12, leading=18, alignment=TA_CENTER, textColor=colors.HexColor("#666666"))
    copyright_style = ParagraphStyle("Copyright", parent=base, fontName=SERIF, fontSize=9, leading=15,
                                     alignment=TA_CENTER, textColor=colors.HexColor("#777777"))
    chap_num_style = ParagraphStyle("ChapterNum", parent=base, fontName=SERIF, fontSize=11, leading=16,
                                    alignment=TA_CENTER, textColor=colors.HexColor("#999999"), spaceAfter=6)
    chap_num_ar_style = ParagraphStyle("ChapterNumAr", parent=chap_num_style, fontName=ARABIC, fontSize=12, leading=18)
    body = ParagraphStyle("Body", parent=base, fontName=SERIF, fontSize=10.5, leading=15.5,
                          alignment=TA_JUSTIFY, firstLineIndent=0.22 * inch, spaceAfter=0)
    body_first = ParagraphStyle("BodyFirst", parent=body, firstLineIndent=0)
    body_rtl = ParagraphStyle("BodyRTL", parent=base, fontName=ARABIC, fontSize=12, leading=20,
                              alignment=TA_RIGHT, firstLineIndent=0, spaceAfter=2)

    def _chapter_title_style(rtl: bool) -> ParagraphStyle:
        return ParagraphStyle("ChapterTitle", parent=base, fontName=(ARABIC_BOLD if rtl else SERIF_BOLD),
                              fontSize=18, leading=24, alignment=TA_CENTER, spaceBefore=4, spaceAfter=26)

    flow: list = []

    # ── Cover image (optional, full-bleed-ish) ──────────────────────────────
    if cover_bytes:
        try:
            from reportlab.lib.utils import ImageReader
            iw, ih = ImageReader(io.BytesIO(cover_bytes)).getSize()
            scale = min(doc.width / iw, doc.height / ih)
            flow += [Image(io.BytesIO(cover_bytes), width=iw * scale, height=ih * scale), PageBreak()]
        except Exception:
            pass

    # ── Title page ──────────────────────────────────────────────────────────
    flow.append(Spacer(1, 2.2 * inch))
    flow.append(Paragraph(_shape(_clean_text(title)) if title_rtl else _html.escape(_clean_text(title), quote=False), title_style))
    flow.append(Spacer(1, 0.3 * inch))
    flow.append(Paragraph(_shape(author) if _is_rtl(author) else _html.escape(author, quote=False), author_style))
    if tagline:
        flow.append(Spacer(1, 0.6 * inch))
        tg = _clean_text(tagline)
        flow.append(Paragraph(_shape(tg) if _is_rtl(tg) else f"“{_html.escape(tg, quote=False)}”", tagline_style))
    flow.append(PageBreak())

    # ── Copyright / colophon page ─────────────────────────────────────────────
    year = datetime.now(timezone.utc).year
    flow.append(Spacer(1, 6.0 * inch))
    flow.append(Paragraph(_html.escape(_clean_text(title), quote=False), copyright_style))
    flow.append(Paragraph(f"© {year} {_html.escape(author, quote=False)}", copyright_style))
    flow.append(Spacer(1, 0.2 * inch))
    flow.append(Paragraph("All rights reserved.", copyright_style))
    flow.append(Paragraph("Produced with G-Ink Studio", copyright_style))
    flow.append(PageBreak())

    # ── Chapters ──────────────────────────────────────────────────────────────
    for ch in chapters:
        rtl = _is_rtl(ch.get("title", "")) or _is_rtl((ch.get("content") or "")[:400])
        raw_title = _clean_text(ch.get("title") or "")
        flow.append(Spacer(1, 0.6 * inch))
        if rtl:
            flow.append(Paragraph(_shape(f"الفصل {_ar_num(ch['number'])}"), chap_num_ar_style))
            outline = f"الفصل {_ar_num(ch['number'])}: {raw_title}"
        else:
            flow.append(Paragraph(f"Chapter {ch['number']}", chap_num_style))
            outline = f"{ch['number']}. {raw_title}"
        ttl = Paragraph(_shape(raw_title) if rtl else _html.escape(raw_title, quote=False), _chapter_title_style(rtl))
        ttl._outline_label = outline  # clean (unshaped) outline label
        flow.append(ttl)
        paras = _paragraphs(ch.get("content") or "")
        scene_break_style = ParagraphStyle(
            "SceneBreak", parent=base, fontName=SERIF, fontSize=10, leading=20,
            alignment=TA_CENTER, spaceBefore=12, spaceAfter=12,
            textColor=colors.HexColor("#999999"),
        )
        for i, para in enumerate(paras):
            if para.strip() == "---":
                flow.append(Paragraph("* * *", scene_break_style))
                continue
            if _is_rtl(para):
                flow.append(Paragraph(_shape(_clean_text(para)), body_rtl))
            else:
                flow.append(Paragraph(_rl_markup(para), body_first if i == 0 else body))
        flow.append(PageBreak())

    doc.build(flow)
    filename = re.sub(r"[^\w\-]", "_", _clean_text(title).lower()) or "story"
    return buf.getvalue(), f"{filename}.pdf", "application/pdf"


# ---------------------------------------------------------------------------
# EPUB Export
# ---------------------------------------------------------------------------

def export_epub(
    title: str,
    author: str,
    tagline: Optional[str],
    genre: Optional[str],
    chapters: list[dict],
    cover_bytes: Optional[bytes] = None,
) -> tuple[bytes, str, str]:
    # Detect the dominant script so e-readers shape + lay out correctly. EPUB
    # readers DO shape Arabic and honor RTL natively, so EPUB is the best format
    # for non-Latin work.
    sample = (title or "") + " " + " ".join((c.get("content") or "")[:200] for c in chapters[:3])
    rtl = _is_rtl(sample)
    lang = "ar" if rtl else "en"

    book = epub.EpubBook()
    book.set_identifier(f"ginink-{re.sub(r'[^a-z0-9]', '-', (title or 'story').lower())}")
    book.set_title(_clean_text(title))
    book.set_language(lang)
    book.add_author(author)
    if rtl:
        book.set_direction("rtl")
    if genre:
        book.add_metadata("DC", "subject", genre)
    if cover_bytes:
        try:
            book.set_cover("cover.jpg", cover_bytes)
        except Exception:
            pass

    css = f"""
        @namespace epub "http://www.idpf.org/2007/ops";
        html, body {{ direction: {'rtl' if rtl else 'ltr'}; }}
        body {{ font-family: {'"Noto Naskh Arabic", ' if rtl else ''}Georgia, "Times New Roman", serif;
               font-size: 1em; line-height: 1.75; margin: 1.4em 1.6em; color: #1a1a1a; }}
        h1 {{ text-align: center; font-size: 1.9em; margin: 1.5em 0 0.3em; line-height: 1.2; }}
        h2 {{ text-align: center; font-size: 1.3em; font-weight: bold; margin: 2.2em 0 1.4em; }}
        .chapnum {{ text-align: center; color: #999; font-variant: small-caps; letter-spacing: .1em;
                    margin-top: 2.4em; font-size: .95em; }}
        p  {{ text-indent: 1.4em; margin: 0; text-align: {'right' if rtl else 'justify'}; }}
        p.first {{ text-indent: 0; }}
        .byline {{ text-align: center; color: #555; font-style: italic; margin-top: .6em; }}
        .tagline {{ text-align: center; font-style: italic; color: #666; margin: 1.6em 2em; }}
        p.scene-break {{ text-indent: 0; text-align: center; color: #aaa; letter-spacing: .25em; margin: 1.2em 0; }}
    """
    style = epub.EpubItem(uid="style", file_name="style/main.css", media_type="text/css", content=css)
    book.add_item(style)

    esc_title = _html.escape(_clean_text(title))
    esc_author = _html.escape(author)
    esc_tagline = _html.escape(_clean_text(tagline)) if tagline else ""

    cover_ch = epub.EpubHtml(title="Title", file_name="title.xhtml", lang=lang)
    cover_ch.content = (
        f"<html><body><h1>{esc_title}</h1>"
        f"<p class='byline'>{esc_author}</p>"
        + (f"<p class='tagline'>{esc_tagline}</p>" if esc_tagline else "")
        + "</body></html>"
    )
    cover_ch.add_item(style)
    book.add_item(cover_ch)

    epub_chapters = [cover_ch]
    toc = []
    for ch in chapters:
        paras = _paragraphs(ch.get("content") or "")
        para_html = ""
        for i, p in enumerate(paras):
            if p.strip() == "---":
                para_html += '<p class="scene-break">* * *</p>\n'
            else:
                para_html += f'<p class="{"first" if i == 0 else ""}">{_epub_markup(p)}</p>\n'
        raw_title = _clean_text(ch.get("title") or "")
        label = f"الفصل {_ar_num(ch['number'])}" if rtl else f"Chapter {ch['number']}"
        nav_title = f"{label}: {raw_title}"
        epub_ch = epub.EpubHtml(title=nav_title, file_name=f"chapter_{ch['number']:03d}.xhtml", lang=lang)
        epub_ch.content = (
            f"<html><body>"
            f"<p class='chapnum'>{_html.escape(label)}</p>"
            f"<h2>{_html.escape(raw_title)}</h2>"
            f"{para_html}</body></html>"
        )
        epub_ch.add_item(style)
        book.add_item(epub_ch)
        epub_chapters.append(epub_ch)
        toc.append(epub.Link(epub_ch.file_name, nav_title, f"ch{ch['number']}"))

    book.toc = toc
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav"] + epub_chapters

    buf = io.BytesIO()
    epub.write_epub(buf, book, {})
    filename = re.sub(r"[^\w\-]", "_", _clean_text(title).lower()) or "story"
    return buf.getvalue(), f"{filename}.epub", "application/epub+zip"


# ---------------------------------------------------------------------------
# DOCX Export (Shunn Standard Manuscript Format)
# ---------------------------------------------------------------------------

def export_docx(
    title: str,
    author: str,
    chapters: list[dict],
) -> tuple[bytes, str, str]:
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    FONT_NAME = "Courier New"
    FONT_SIZE = Pt(12)

    def _set_font(run):
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE

    def _para(text: str, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.line_spacing = Pt(24)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        _set_font(run)
        run.bold = bold
        return p

    total_words = sum(_word_count(ch["content"]) for ch in chapters)
    word_est = f"~{round(total_words / 1000)}k words" if total_words > 999 else f"{total_words} words"

    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run(f"{author}\n{word_est}")
    _set_font(run)

    doc.add_paragraph()
    for _ in range(8):
        _para("")
    _para(_clean_text(title).upper(), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _para("by", align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(author, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    for ch in chapters:
        ch_title = f"Chapter {ch['number']}: {_clean_text(ch['title'])}"
        for _ in range(10):
            _para("")
        _para(ch_title.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        doc.add_paragraph()

        first = True
        for para_text in _paragraphs(ch["content"]):
            if para_text.strip() == "---":
                sb = doc.add_paragraph()
                sb.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sb.paragraph_format.line_spacing = Pt(24)
                sb.paragraph_format.space_before = Pt(0)
                sb.paragraph_format.space_after = Pt(0)
                run = sb.add_run("* * *")
                _set_font(run)
                first = False
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = Pt(24)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Inches(0) if first else Inches(0.5)
            run = p.add_run(_clean_text(para_text))
            _set_font(run)
            first = False
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    filename = re.sub(r"[^\w\-]", "_", _clean_text(title).lower()) or "story"
    return buf.getvalue(), f"{filename}_manuscript.docx", \
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ---------------------------------------------------------------------------
# Submission Package (zip: PDF + EPUB + DOCX + synopsis TXT)
# ---------------------------------------------------------------------------

async def export_submission_package(
    title: str,
    author: str,
    tagline: Optional[str],
    genre: Optional[str],
    chapters: list[dict],
    synopsis: Optional[str] = None,
) -> tuple[bytes, str, str]:
    pdf_bytes, pdf_name, _ = export_pdf(title, author, tagline, chapters)
    epub_bytes, epub_name, _ = export_epub(title, author, tagline, genre, chapters)
    docx_bytes, docx_name, _ = export_docx(title, author, chapters)

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(pdf_name, pdf_bytes)
        zf.writestr(epub_name, epub_bytes)
        zf.writestr(docx_name, docx_bytes)
        if synopsis:
            zf.writestr("synopsis.txt", synopsis.encode("utf-8"))
        readme = (
            f"Submission Package — {title}\n"
            f"Author: {author}\n"
            f"Genre: {genre or 'unspecified'}\n"
            f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}\n\n"
            f"Files included:\n"
            f"  {pdf_name}  — reading copy\n"
            f"  {epub_name} — e-reader format\n"
            f"  {docx_name} — Shunn standard manuscript\n"
            + ("  synopsis.txt     — one-page synopsis\n" if synopsis else "")
        )
        zf.writestr("README.txt", readme.encode("utf-8"))

    filename = re.sub(r"[^\w\-]", "_", _clean_text(title).lower()) or "story"
    return buf.getvalue(), f"{filename}_submission.zip", "application/zip"
